import os
import argparse
import sys
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

def sanitize_filename_part(text):
    """清理文件名或页眉中可能导致语法错误的非法字符"""
    return re.sub(r'[\/*?:"<>|]', "", text)

def add_page_number(run):
    """在指定的 run 中插入 Word 自动页码字段"""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def create_code_docx(input_file, output_docx, software_name, version):
    # 规范化路径
    input_file = os.path.normpath(input_file)
    output_docx = os.path.normpath(output_docx)
    abs_output = os.path.abspath(output_docx)
    
    if not os.path.exists(os.path.dirname(abs_output)):
        os.makedirs(os.path.dirname(abs_output), exist_ok=True)

    # 创建文档
    doc = Document()
    
    # 【核心修复】彻底清除文档初始化时自带的所有段落
    for p in doc.paragraphs:
        p._element.getparent().remove(p._element)

    # 1. 页面设置
    section = doc.sections[0]
    section.top_margin = Inches(0.4)
    section.bottom_margin = Inches(0.4)
    section.left_margin = Inches(0.8) 
    section.right_margin = Inches(0.4)
    
    # 【强制行号修正逻辑】
    sectPr = section._sectPr
    # 移除现有的行号设置
    for ln in sectPr.xpath('w:lnNumType'):
        sectPr.remove(ln)
        
    lnNumType = OxmlElement('w:lnNumType')
    lnNumType.set(qn('w:countBy'), '1')
    # 【关键修改】如果设置 1 出来的是 2，那么我们将其设为 0，这样第一行就会显示为 1
    lnNumType.set(qn('w:start'), '0') 
    lnNumType.set(qn('w:restart'), 'continuous')
    sectPr.append(lnNumType)
    
    # 【禁用文档网格，防止字符换行到网格导致长行被拆分】
    for pgMar in sectPr.xpath('w:pgMar'):
        pgMar.set(qn('w:docGrid'), '0')
    
    # 设置文档网格类型为只在首页显示
    docGrid = OxmlElement('w:docGrid')
    docGrid.set(qn('w:type'), 'default')
    docGrid.set(qn('w:linePitch'), '312')
    docGrid.set(qn('w:charPitch'), '240')  # 禁用字符网格
    docGrid.set(qn('w:layoutMode'), 'snapToGrid')
    # 移除可能导致换行的设置
    docGrid.set(qn('w:noEndnote'), '0')
    docGrid.set(qn('w:rtlGutter'), '0')
    docGrid.set(qn('w:view'), 'normal')
    docGrid.set(qn('w:zoom'), '100')
    docGrid.set(qn('w:recentSpacesChange'), '0')
    
    # 确保pgMar有docGrid属性
    pgMar_list = sectPr.xpath('w:pgMar')
    if pgMar_list:
        pgMar_list[0].set(qn('w:docGrid'), '0')
        pgMar_list[0].set(qn('w:gridType'), 'none')  # 禁用网格换行

    # 2. 定义全局样式
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(9.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    pf.line_spacing = Pt(13.9) 
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)

    # 3. 页眉页脚设置
    safe_name = sanitize_filename_part(software_name)
    safe_version = sanitize_filename_part(version)
    
    header = section.header
    p_header = header.paragraphs[0]
    p_header.clear()
    
    # 添加制表符（用于居中和右对齐）
    tab_stops = p_header.paragraph_format.tab_stops
    tab_stops.add_tab_stop(Inches(3.0), WD_ALIGN_PARAGRAPH.CENTER)
    tab_stops.add_tab_stop(Inches(6.0), WD_ALIGN_PARAGRAPH.RIGHT)
    
    # 添加软件名称和版本号（居中）
    p_header.add_run('\t')
    run_header = p_header.add_run(f"{safe_name} {safe_version} 源代码")
    run_header.font.size = Pt(9)
    
    # 添加页码（右对齐）
    p_header.add_run('\t')
    run_page = p_header.add_run()
    run_page.font.size = Pt(9)
    add_page_number(run_page)

    # 清空页脚，不再在页脚显示页码
    footer = section.footer
    p_footer = footer.paragraphs[0]
    p_footer.clear()

    # 4. 读取内容并严格过滤
    if not os.path.exists(input_file):
        print(f"错误：找不到输入文件 {input_file}")
        return

    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    processed_lines = []
    for line in lines:
        clean = line.rstrip('\n\r')
        # 过滤 Markdown 代码块标记
        if clean.strip().startswith('```'):
            continue
        # 过滤掉所有空行（包括内容前后的）
        if not clean.strip():
            continue
        # 拆分超过80字符的长行，避免Word自动换行导致行数不准
        if len(clean) > 80:
            # 在合适位置（逗号、括号、空格）拆分
            parts = []
            while len(clean) > 80:
                # 找最后一个能在80字符内的逗号或括号
                split_pos = min(80, len(clean))
                for i in range(min(80, len(clean)), 0, -1):
                    if clean[i-1] in [',', '(', ')', '>', '<', '=']:
                        split_pos = i
                        break
                parts.append(clean[:split_pos])
                clean = clean[split_pos:]
            if clean:
                parts.append(clean)
            processed_lines.extend(parts)
        else:
            processed_lines.append(clean)

    # 【截断逻辑：在拆分长行后，如果总行数超过3000，截断】
    MAX_PARAGRAPHS = 3000
    if len(processed_lines) > MAX_PARAGRAPHS:
        # 保留前一半和后一半，但总行数不超过MAX_PARAGRAPHS
        half = MAX_PARAGRAPHS // 2
        processed_lines = processed_lines[:half] + processed_lines[-half:]
        print(f"⚠️ 代码较长，已截断为前{half}行和后{half}行")

    # 5. 写入内容
    print(f"原始行数：{len(lines)}，写入行数：{len(processed_lines)} 行")

    paragraph_count = 0

    for content in processed_lines:
        p = doc.add_paragraph()
        paragraph_count += 1
        # 强制应用固定行距 13.9
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        p.paragraph_format.line_spacing = Pt(13.9)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

        # 添加文字，设置不换行
        run = p.add_run(content)
        # 设置 run 级别的 wordWrap 为 false，禁用自动换行
        rPr = run._r.get_or_add_rPr()
        # 禁用断字
        noHyphen = OxmlElement('w:noHyphenation')
        noHyphen.set(qn('w:val'), '1')
        rPr.append(noHyphen)
        # 设置溢出标点不换行
        overflowPunct = OxmlElement('w:overflowPunct')
        overflowPunct.set(qn('w:val'), '1')
        rPr.append(overflowPunct)
        # 设置字符不压缩
        compressPunct = OxmlElement('w:compressPunct')
        compressPunct.set(qn('w:val'), '0')
        rPr.append(compressPunct)
        # 设置字符间距不调整
        autoSpaceDE = OxmlElement('w:autoSpaceDE')
        autoSpaceDE.set(qn('w:val'), '0')
        rPr.append(autoSpaceDE)
        # 设置字符水平缩放为 100%（不压缩）
        charSpacing = OxmlElement('w:charSpacing')
        charSpacing.set(qn('w:val'), '0')
        rPr.append(charSpacing)

        # 段落级别禁用换行
        pPr = p._p.get_or_add_pPr()
        # 禁止在标点处换行
        overflowPunct2 = OxmlElement('w:overflowPunct')
        overflowPunct2.set(qn('w:val'), '1')
        pPr.append(overflowPunct2)

    total_written = paragraph_count

    try:
        doc.save(abs_output)
        print("-" * 50)
        print(f"✅ 转换成功！")
        print(f"📝 行间距: 13.9 磅 (固定值)")
        print(f"🔢 行号起始补偿: 已设为 0 (预期首行显示为 1)")
        print(f"📂 存储路径: {abs_output}")
        print(f"📊 总段落数: {total_written}（约{total_written//50}页）")
        print("-" * 50)
    except PermissionError:
        print(f"❌ 错误：文件 '{abs_output}' 正在被 Word 使用，请关闭后重试。")
    except Exception as e:
        print(f"发生未知错误: {e}")

def code_convert_main(input_md, output_docx, software_name="未知软件", version="v1.0"):
    # 直接调用你写好的核心函数
    create_code_docx(input_md, output_docx, software_name, version)
    
if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('--input_md', required=True)
    parser.add_argument('--output_docx', required=True)
    parser.add_argument('--software_name', required=True)
    parser.add_argument('--version', required=True)
    args = parser.parse_args()
    create_code_docx(args.input_md, args.output_docx, args.software_name, args.version)