#!/usr/bin/env python3
import argparse
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


def set_font(run, name="宋体", size=10.5, bold=False):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, end])
    set_font(run, size=9)


def summary_convert_main(input_md, output_docx, software_name="未知软件", version="V1.0"):
    markdown = open(input_md, "r", encoding="utf-8").read()
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.4)
    section.right_margin = Cm(2.4)
    section.header_distance = Cm(1.0)
    section.footer_distance = Cm(1.0)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(header.add_run(f"{software_name} {version} 申请信息摘要"), size=9)
    add_page_number(section.footer.paragraphs[0])

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(title.add_run("计算机软件著作权登记信息采集表"), size=18, bold=True)

    for raw_line in markdown.splitlines():
        line = raw_line.strip()
        if not line or line.startswith("### 计算机软件著作权登记信息采集表"):
            continue
        heading = re.match(r"^(#{1,6})\s+(.+)$", line)
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_after = Pt(4)
        if heading:
            set_font(paragraph.add_run(heading.group(2)), size=12, bold=True)
        else:
            clean = re.sub(r"^[-*]\s+", "", line)
            clean = clean.replace("**", "").replace("`", "")
            set_font(paragraph.add_run(clean), size=10.5)
    doc.save(output_docx)


if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--input_md", required=True)
    parser.add_argument("--output_docx", required=True)
    parser.add_argument("--software_name", default="未知软件")
    parser.add_argument("--version", default="V1.0")
    args = parser.parse_args()
    summary_convert_main(args.input_md, args.output_docx, args.software_name, args.version)
