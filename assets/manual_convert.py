#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Markdown to Word Converter for Software Copyright Application
将Markdown用户手册转换为符合软著申请要求的Word文档
"""

import argparse
import re
import os
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_TAB_ALIGNMENT
    from docx.oxml.shared import OxmlElement, qn
    import markdown
    from bs4 import BeautifulSoup
except ImportError as e:
    print(f"错误：缺少必要的依赖库")
    print(f"请安装：pip install python-docx markdown beautifulsoup4")
    print(f"详细错误：{e}")
    sys.exit(1)


class MarkdownToWordConverter:
    """Markdown转Word转换器"""
    
    def __init__(self, software_name, version, cover_template=None):
        self.software_name = software_name
        self.version = version
        self.cover_template = cover_template
        self.doc = None
        
    def create_from_cover(self):
        """从封面模板创建文档"""
        if self.cover_template and os.path.exists(self.cover_template):
            # 加载封面模板
            self.doc = Document(self.cover_template)
            # 替换正文中的变量（在 run 级别操作，保证不破坏原有字号和样式，例如 36 号字体）
            for paragraph in self.doc.paragraphs:
                for run in paragraph.runs:
                    if not run.text:
                        continue
                    text = run.text
                    # 各种占位符写法统一替换
                    text = text.replace('{{软件名称}}', self.software_name)
                    text = text.replace('{{版本号}}', self.version)
                    text = text.replace('{软件名称}', self.software_name)
                    text = text.replace('{版本号}', self.version)
                    text = text.replace('软件名称', self.software_name)
                    text = text.replace('版本号', self.version)
                    # 去掉普通文本里的花括号，防止 {{软件名称}} 残留一个 {
                    text = text.replace('{', '').replace('}', '')
                    run.text = text
                    # 保留模板自身的字号和样式
                    self._apply_formatting(run, skip_format=True)
                  
            # 替换正文表格中的变量（如果有），同样在 run 级别处理
            for table in self.doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if not run.text:
                                    continue
                                text = run.text
                                text = text.replace('{{软件名称}}', self.software_name)
                                text = text.replace('{{版本号}}', self.version)
                                text = text.replace('{软件名称}', self.software_name)
                                text = text.replace('{版本号}', self.version)
                                text = text.replace('软件名称', self.software_name)
                                text = text.replace('版本号', self.version)
                                text = text.replace('{', '').replace('}', '')
                                run.text = text
                                self._apply_formatting(run, skip_format=True)

            # 替换页眉和页脚中的变量（如果有），同时完全保留模板中的样式和页码域
            for section in self.doc.sections:
                # 处理各类页眉：默认 / 首页 / 偶数页
                header_attrs = ["header", "first_page_header", "even_page_header"]
                for attr in header_attrs:
                    header = getattr(section, attr, None)
                    if header is None:
                        continue
                    # 直接在 run 级别做文本替换，避免破坏 {PAGE} 等域代码
                    for paragraph in header.paragraphs:
                        for run in paragraph.runs:
                            if not run.text:
                                continue
                            text = run.text
                            # 支持多种占位符写法，并确保最终不残留花括号
                            text = text.replace('{{软件名称}}', self.software_name)
                            text = text.replace('{{版本号}}', self.version)
                            text = text.replace('{软件名称}', self.software_name)
                            text = text.replace('{版本号}', self.version)
                            # 兜底：直接替换关键字本身，防止因为空格等原因未命中
                            text = text.replace('软件名称', self.software_name)
                            text = text.replace('版本号', self.version)
                            # 去掉可能残留的花括号，但不会影响 Word 的 PAGE 域（域的大括号不是普通字符）
                            text = text.replace('{', '').replace('}', '')
                            run.text = text
                            self._apply_formatting(run, skip_format=True)
                    # 页眉中的表格单元格
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if not run.text:
                                            continue
                                        text = run.text
                                        text = text.replace('{{软件名称}}', self.software_name)
                                        text = text.replace('{{版本号}}', self.version)
                                        text = text.replace('{软件名称}', self.software_name)
                                        text = text.replace('{版本号}', self.version)
                                        text = text.replace('软件名称', self.software_name)
                                        text = text.replace('版本号', self.version)
                                        text = text.replace('{', '').replace('}', '')
                                        run.text = text
                                        self._apply_formatting(run, skip_format=True)

                # 处理各类页脚：默认 / 首页 / 偶数页（一般很少放软件名称，但一并支持）
                footer_attrs = ["footer", "first_page_footer", "even_page_footer"]
                for attr in footer_attrs:
                    footer = getattr(section, attr, None)
                    if footer is None:
                        continue
                    for paragraph in footer.paragraphs:
                        for run in paragraph.runs:
                            if not run.text:
                                continue
                            text = run.text
                            text = text.replace('{{软件名称}}', self.software_name)
                            text = text.replace('{{版本号}}', self.version)
                            text = text.replace('{软件名称}', self.software_name)
                            text = text.replace('{版本号}', self.version)
                            text = text.replace('软件名称', self.software_name)
                            text = text.replace('版本号', self.version)
                            text = text.replace('{', '').replace('}', '')
                            run.text = text
                            self._apply_formatting(run, skip_format=True)
                    for table in footer.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        if not run.text:
                                            continue
                                        text = run.text
                                        text = text.replace('{{软件名称}}', self.software_name)
                                        text = text.replace('{{版本号}}', self.version)
                                        text = text.replace('{软件名称}', self.software_name)
                                        text = text.replace('{版本号}', self.version)
                                        text = text.replace('软件名称', self.software_name)
                                        text = text.replace('版本号', self.version)
                                        text = text.replace('{', '').replace('}', '')
                                        run.text = text
                                        self._apply_formatting(run, skip_format=True)
        else:
            # 创建新文档并添加简单封面
            self.doc = Document()
            self._add_simple_cover()
    
    def _add_simple_cover(self):
        """添加简单封面"""
        # 添加多个空段落
        for _ in range(5):
            self.doc.add_paragraph()
        
        # 软件名称
        title = self.doc.add_paragraph()
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = title.add_run("软件用户说明书")
        run.font.size = Pt(22)
        run.font.bold = True
        run.font.name = '宋体'
        
        self.doc.add_paragraph()
        
        # 软件名称和版本号
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = name_para.add_run(f"软件名称：{self.software_name}")
        run.font.size = Pt(16)
        run.font.name = '宋体'
        
        version_para = self.doc.add_paragraph()
        version_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = version_para.add_run(f"版本号：{self.version}")
        run.font.size = Pt(16)
        run.font.name = '宋体'
        
        # 添加分页符
        self.doc.add_page_break()
    
    def parse_markdown(self, markdown_file):
        """解析Markdown文件"""
        with open(markdown_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        
        # 转换为HTML
        html = markdown.markdown(
            md_content,
            extensions=['tables', 'fenced_code', 'codehilite']
        )
        
        return html
    
    def html_to_word(self, html):
        """将HTML转换为Word内容"""
        soup = BeautifulSoup(html, 'html.parser')
        
        # 处理每个HTML元素
        for element in soup.find_all(True):
            tag_name = element.name
            
            if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                self._add_heading(element, tag_name)
            elif tag_name == 'p':
                self._add_paragraph(element)
            elif tag_name == 'ul':
                self._add_list(element, ordered=False)
            elif tag_name == 'ol':
                self._add_list(element, ordered=True)
            elif tag_name == 'table':
                self._add_table(element)
            elif tag_name == 'pre':
                self._add_code_block(element)
            elif tag_name == 'blockquote':
                self._add_blockquote(element)
            elif tag_name == 'hr':
                self._add_horizontal_rule()
    
    def _add_heading(self, element, level):
        """添加标题"""
        text = element.get_text().strip()
        if not text:
            return
        
        # 映射标题级别
        level_map = {'h1': 1, 'h2': 2, 'h3': 3, 'h4': 4, 'h5': 5, 'h6': 6}
        heading_level = level_map.get(level, 1)
        
        heading = self.doc.add_heading(text, level=heading_level)
        self._apply_formatting(heading.runs[0])
    
    def _add_paragraph(self, element):
        """添加段落"""
        text = element.get_text().strip()
        if not text:
            return
        
        para = self.doc.add_paragraph(text)
        self._apply_formatting(para.runs[0])
    
    def _add_list(self, element, ordered=False):
        """添加列表"""
        items = element.find_all('li', recursive=False)
        
        for idx, item in enumerate(items, start=1):
            text = item.get_text().strip()
            if not text:
                continue
            
            # 优先尝试使用内置列表样式；若当前文档（或封面模板）中不存在该样式，则降级为普通段落并手动添加符号
            try:
                if ordered:
                    para = self.doc.add_paragraph(text, style='List Number')
                else:
                    para = self.doc.add_paragraph(text, style='List Bullet')
            except KeyError:
                # 某些自定义模板中可能没有 List Bullet / List Number 样式
                if ordered:
                    display_text = f"{idx}. {text}"
                else:
                    display_text = f"• {text}"
                para = self.doc.add_paragraph(display_text)
            
            if para.runs:
                self._apply_formatting(para.runs[0])
    
    def _add_table(self, element):
        """添加表格"""
        rows = element.find_all('tr')
        if not rows:
            return
        
        # 获取列数
        cols = len(rows[0].find_all(['th', 'td'], recursive=False))
        
        # 创建表格
        table = self.doc.add_table(rows=len(rows), cols=cols)
        table.style = 'Table Grid'
        
        # 填充表格内容
        for row_idx, row in enumerate(rows):
            cells = row.find_all(['th', 'td'], recursive=False)
            for col_idx, cell in enumerate(cells):
                if col_idx < cols:
                    text = cell.get_text().strip()
                    table.rows[row_idx].cells[col_idx].text = text
                    # 表头加粗
                    if cell.name == 'th':
                        for paragraph in table.rows[row_idx].cells[col_idx].paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
    
    def _add_code_block(self, element):
        """添加代码块"""
        code_text = element.get_text()
        if not code_text:
            return
        
        para = self.doc.add_paragraph(code_text)
        # 尝试设置样式，如果不存在则跳过
        try:
            para.style = 'No Spacing'
        except KeyError:
            pass  # 样式不存在时使用默认样式
        if para.runs:
            run = para.runs[0]
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(64, 64, 64)
    
    def _add_blockquote(self, element):
        """添加引用块"""
        text = element.get_text().strip()
        if not text:
            return
        
        para = self.doc.add_paragraph(text)
        if para.runs:
            run = para.runs[0]
            run.font.italic = True
            # 设置左缩进
            para.paragraph_format.left_indent = Inches(0.5)
    
    def _add_horizontal_rule(self):
        """添加水平线"""
        para = self.doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = para.add_run('─' * 50)
    
    def _apply_formatting(self, run, skip_format=False):
        """应用格式设置
        
        :param run: 需要应用格式的 run 对象
        :param skip_format: 为 True 时不修改任何格式（用于封面模板，保留原样式）
        """
        # 封面模板中的内容需要保留原有样式，直接跳过
        if skip_format or run is None:
            return
        
        # 设置字体
        run.font.name = '宋体'
        run.font.size = Pt(10.5)
        
        # 设置段落格式
        paragraph = run._element.getparent()
        if paragraph is not None:
            # 设置行距为单倍行距（确保每页30行以上）
            pPr = paragraph.get_or_add_pPr()
            spacing = OxmlElement('w:spacing')
            spacing.set(qn('w:line'), '240')  # 单倍行距
            spacing.set(qn('w:lineRule'), 'auto')
            pPr.append(spacing)
            
            # 设置段前段后间距
            spacing.set(qn('w:before'), '40')
            spacing.set(qn('w:after'), '40')
    
    def add_header_with_page_number(self):
        """使用与源代码文档相同的单段页眉，避开模板中的浮动 PAGE 内容控件。"""
        for section in self.doc.sections:
            # 模板历史页眉把 PAGE 域放在带 framePr 的 SDT 中，标题另占一段。
            # Word 会把两者视觉叠放，LibreOffice 则按两个段落导出，导致 PAGE 1
            # 独占一行且缓存页码不更新。因此必须清空 header 根节点，而不只是
            # paragraph.clear()（后者无法删除包在顶层 SDT 中的字段）。
            section.different_first_page_header_footer = False

            # 需要处理的所有可能页眉：默认 / 首页 / 偶数页
            headers = [section.header]
            # 某些 python-docx 版本可能不存在这些属性，这里做兼容处理
            for attr in ("first_page_header", "even_page_header"):
                try:
                    headers.append(getattr(section, attr))
                except AttributeError:
                    continue

            for header in headers:
                if header is None:
                    continue

                header_element = header._element
                for child in list(header_element):
                    header_element.remove(child)
                header_para = header.add_paragraph()
                header_para.paragraph_format.space_before = Pt(0)
                header_para.paragraph_format.space_after = Pt(0)

                # 页眉制表位必须落在当前 section 的可打印宽度内。
                # 用户手册模板左右页边距各 1.25 英寸，可用宽度约 5.77 英寸；
                # 旧实现把页码固定放在 6.0 英寸处，生产 LibreOffice 会直接裁掉。
                usable_width = section.page_width - section.left_margin - section.right_margin

                # 添加制表符（标题居中、页码贴右）
                tab_stops = header_para.paragraph_format.tab_stops
                tab_stops.add_tab_stop(usable_width // 2, WD_TAB_ALIGNMENT.CENTER)
                tab_stops.add_tab_stop(usable_width, WD_TAB_ALIGNMENT.RIGHT)

                # 添加软件名称和版本号（居中）
                header_para.add_run('\t')
                run_name = header_para.add_run(f'{self.software_name} {self.version} 用户手册')
                run_name.font.size = Pt(9)
                run_name.font.name = '宋体'

                # 与源代码文档保持一致：PAGE 域和标题位于同一普通段落。
                header_para.add_run('\t')
                run_page = header_para.add_run()
                run_page.font.size = Pt(9)
                run_page.font.name = '宋体'

                # 插入页码域 { PAGE }
                fldChar1 = OxmlElement('w:fldChar')
                fldChar1.set(qn('w:fldCharType'), 'begin')

                instrText = OxmlElement('w:instrText')
                instrText.set(qn('xml:space'), 'preserve')
                instrText.text = "PAGE"

                fldChar2 = OxmlElement('w:fldChar')
                fldChar2.set(qn('w:fldCharType'), 'end')

                run_page._r.append(fldChar1)
                run_page._r.append(instrText)
                run_page._r.append(fldChar2)
    
    def set_different_first_page(self):
        """设置首页不同（封面不带页眉）"""
        section = self.doc.sections[0]
        section.different_first_page_header_footer = True
    
    def convert(self, input_md, output_docx):
        """执行转换"""
        # 1. 从封面创建文档
        self.create_from_cover()
        
        # 2. 解析Markdown
        html = self.parse_markdown(input_md)
        
        # 3. 转换HTML到Word
        self.html_to_word(html)
        
        # 4. 无论是否使用封面模板，都规范化为与源代码文档相同的单段页眉。
        self.add_header_with_page_number()
        
        # 6. 保存文档
        self.doc.save(output_docx)
        print(f"转换成功！输出文件：{output_docx}")


def main():
    """主函数"""
    parser = argparse.ArgumentParser(
        description='将Markdown用户手册转换为符合软著申请要求的Word文档'
    )
    
    parser.add_argument(
        '--input_md',
        required=True,
        help='输入的Markdown文件路径'
    )
    
    parser.add_argument(
        '--output_docx',
        required=True,
        help='输出的docx文件路径'
    )
    
    parser.add_argument(
        '--software_name',
        required=True,
        help='软件名称'
    )
    
    parser.add_argument(
        '--version',
        required=True,
        help='版本号'
    )
    
    parser.add_argument(
        '--cover',
        help='封面模板docx文件路径（可选）'
    )
    
    args = parser.parse_args()
    
    # 验证输入文件
    if not os.path.exists(args.input_md):
        print(f"错误：输入文件不存在：{args.input_md}")
        sys.exit(1)
    
    # 验证封面模板（如果提供）
    if args.cover and not os.path.exists(args.cover):
        print(f"警告：封面模板不存在：{args.cover}，将使用默认封面")
        args.cover = None
    
    # 创建转换器并执行转换
    converter = MarkdownToWordConverter(
        software_name=args.software_name,
        version=args.version,
        cover_template=args.cover
    )
    
    try:
        converter.convert(args.input_md, args.output_docx)
    except Exception as e:
        print(f"转换失败：{e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)
        
# 在 manual_convert.py 底部添加这个函数
def manual_convert_main(input_md, output_docx, software_name="未知软件", version="v1.0", cover=None):
    converter = MarkdownToWordConverter(
        software_name=software_name,
        version=version,
        cover_template=cover
    )
    converter.convert(input_md, output_docx)


if __name__ == '__main__':
    main()
